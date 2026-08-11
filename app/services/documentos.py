"""Geração de documentos a partir do texto (Markdown) produzido pelos agentes.

O "Gabinete Designer": exporta a peça/notificação/parecer para Word (.docx) e
PDF, e tabelas para Excel (.xlsx). Tudo pure-Python (seguro no Docker), com o
cabeçalho do escritório. Não depende do Google.
"""

import io
import re
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font as XLFont

_COR_PADRAO = "#9A6A3A"


@dataclass
class Timbre:
    """Identidade visual aplicada ao cabeçalho/rodapé das peças exportadas."""

    nome: str = ""
    subtitulo: str = ""
    cor: str = _COR_PADRAO
    rodape: str = ""
    logo: bytes | None = None  # PNG/JPEG já decodificado (opcional)


def _hex_rgb(cor: str) -> tuple[int, int, int]:
    """Converte '#RRGGBB' em (r, g, b); cai no padrão se inválido."""
    c = (cor or "").lstrip("#")
    if len(c) != 6:
        c = _COR_PADRAO.lstrip("#")
    try:
        return int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
    except ValueError:
        c = _COR_PADRAO.lstrip("#")
        return int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)

# ── Parser de Markdown → blocos ─────────────────────────────────

_BULLET = re.compile(r"^[-*]\s+")
_NUM = re.compile(r"^\d+\.\s+")
_SEP = re.compile(r"^:?-{2,}:?$")


def parse_markdown(texto: str) -> list[tuple[str, Any]]:
    """Converte o Markdown em blocos: (h1/h2/h3/paragrafo/lista/tabela, conteúdo)."""
    linhas = texto.replace("\r\n", "\n").split("\n")
    blocos: list[tuple[str, Any]] = []
    par: list[str] = []
    lista: list[str] = []

    def flush_par() -> None:
        if par:
            blocos.append(("paragrafo", " ".join(par).strip()))
            par.clear()

    def flush_lista() -> None:
        if lista:
            blocos.append(("lista", list(lista)))
            lista.clear()

    i = 0
    while i < len(linhas):
        bruta = linhas[i]
        ln = bruta.strip()
        if not ln:
            flush_par()
            flush_lista()
        elif ln.startswith("### "):
            flush_par(), flush_lista(), blocos.append(("h3", ln[4:].strip()))
        elif ln.startswith("## "):
            flush_par(), flush_lista(), blocos.append(("h2", ln[3:].strip()))
        elif ln.startswith("# "):
            flush_par(), flush_lista(), blocos.append(("h1", ln[2:].strip()))
        elif _BULLET.match(ln):
            flush_par()
            lista.append(_BULLET.sub("", ln).strip())
        elif _NUM.match(ln):
            flush_par()
            lista.append(_NUM.sub("", ln).strip())
        elif ln.startswith("|") and ln.endswith("|"):
            flush_par()
            flush_lista()
            tabela: list[list[str]] = []
            while i < len(linhas) and linhas[i].strip().startswith("|"):
                celulas = [c.strip() for c in linhas[i].strip().strip("|").split("|")]
                if not all(_SEP.match(c) for c in celulas if c):
                    tabela.append([_sem_negrito(c) for c in celulas])
                i += 1
            blocos.append(("tabela", tabela))
            continue
        else:
            flush_lista()
            par.append(ln)
        i += 1
    flush_par()
    flush_lista()
    return blocos


def _sem_negrito(texto: str) -> str:
    return texto.replace("**", "")


# ── Word (.docx) ────────────────────────────────────────────────


def _runs_negrito(paragrafo: Any, texto: str) -> None:
    """Escreve o texto no parágrafo tratando **negrito** inline."""
    for i, parte in enumerate(texto.split("**")):
        if parte:
            run = paragrafo.add_run(parte)
            run.bold = i % 2 == 1


def gerar_docx(
    titulo: str, markdown: str, escritorio: str = "", timbre: Timbre | None = None
) -> bytes:
    """Gera um .docx com o timbre do escritório e o conteúdo formatado."""
    t = timbre or Timbre(nome=escritorio)
    cor = RGBColor(*_hex_rgb(t.cor))
    doc = Document()
    if t.logo:
        try:
            doc.add_picture(io.BytesIO(t.logo), width=Inches(1.6))
        except Exception:  # noqa: BLE001 - logo inválido não impede a exportação
            pass
    if t.nome:
        cab = doc.add_paragraph()
        r = cab.add_run(t.nome)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = cor
    if t.subtitulo:
        sub = doc.add_paragraph()
        rs = sub.add_run(t.subtitulo)
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    if t.rodape:
        rod = doc.sections[0].footer.paragraphs[0]
        rr = rod.add_run(t.rodape)
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if titulo:
        h = doc.add_heading(titulo, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for tipo, conteudo in parse_markdown(markdown):
        if tipo in ("h1", "h2", "h3"):
            doc.add_heading(conteudo, level=int(tipo[1]))
        elif tipo == "paragrafo":
            _runs_negrito(doc.add_paragraph(), conteudo)
        elif tipo == "lista":
            for item in conteudo:
                _runs_negrito(doc.add_paragraph(style="List Bullet"), item)
        elif tipo == "tabela" and conteudo:
            n_col = max(len(linha) for linha in conteudo)
            tabela = doc.add_table(rows=0, cols=n_col)
            tabela.style = "Light Grid Accent 1"
            for r_idx, linha in enumerate(conteudo):
                celulas = tabela.add_row().cells
                for c_idx in range(n_col):
                    txt = linha[c_idx] if c_idx < len(linha) else ""
                    celulas[c_idx].text = txt
                    if r_idx == 0:
                        for p in celulas[c_idx].paragraphs:
                            for run in p.runs:
                                run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── PDF ─────────────────────────────────────────────────────────

_SUBS_PDF = {
    "—": "-",
    "–": "-",
    "“": '"',
    "”": '"',
    "‘": "'",
    "’": "'",
    "•": "-",
    "…": "...",
    "→": "->",
    "✓": "[ok]",
    "₂": "2",
    "º": "o",
    "ª": "a",
}


def _lat1(texto: str) -> str:
    """Torna o texto seguro para as fontes core do fpdf (latin-1)."""
    for k, v in _SUBS_PDF.items():
        texto = texto.replace(k, v)
    return texto.encode("latin-1", "replace").decode("latin-1")


class _PDF(FPDF):
    timbre = Timbre()

    def header(self) -> None:  # noqa: D102 - chamado pelo fpdf a cada página
        t = self.timbre
        if not (t.nome or t.logo):
            return
        r, g, b = _hex_rgb(t.cor)
        x_texto = self.l_margin
        if t.logo:
            try:
                self.image(io.BytesIO(t.logo), x=self.l_margin, y=9, h=13)
                x_texto = self.l_margin + 36
            except Exception:  # noqa: BLE001 - logo inválido não impede o PDF
                pass
        self.set_xy(x_texto, 10)
        self.set_font("Helvetica", "B", 12)
        self.set_text_color(r, g, b)
        self.multi_cell(0, 7, _lat1(t.nome))
        if t.subtitulo:
            self.set_x(x_texto)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(0x60, 0x60, 0x60)
            self.multi_cell(0, 5, _lat1(t.subtitulo))
        self.set_draw_color(r, g, b)
        self.line(self.l_margin, self.get_y() + 1, self.w - self.r_margin, self.get_y() + 1)
        self.ln(5)
        self.set_text_color(0x1D, 0x22, 0x30)

    def footer(self) -> None:  # noqa: D102 - chamado pelo fpdf a cada página
        if not self.timbre.rodape:
            return
        self.set_y(-14)
        self.set_font("Helvetica", "", 7)
        self.set_text_color(0x80, 0x80, 0x80)
        self.cell(0, 6, _lat1(self.timbre.rodape), align="C")


def gerar_pdf(
    titulo: str, markdown: str, escritorio: str = "", timbre: Timbre | None = None
) -> bytes:
    """Gera um PDF com o timbre do escritório e o conteúdo formatado."""
    pdf = _PDF(format="A4")
    pdf.timbre = timbre or Timbre(nome=escritorio)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(20, 18, 20)
    pdf.add_page()

    if titulo:
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 9, _lat1(titulo))
        pdf.ln(2)

    for tipo, conteudo in parse_markdown(markdown):
        pdf.set_x(pdf.l_margin)
        if tipo in ("h1", "h2", "h3"):
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", {"h1": 14, "h2": 12, "h3": 11}[tipo])
            pdf.multi_cell(0, 7, _lat1(_sem_negrito(conteudo)))
            pdf.ln(1)
        elif tipo == "paragrafo":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, _lat1(_sem_negrito(conteudo)))
            pdf.ln(1)
        elif tipo == "lista":
            pdf.set_font("Helvetica", "", 11)
            for item in conteudo:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, _lat1("-  " + _sem_negrito(item)))
            pdf.ln(1)
        elif tipo == "tabela" and conteudo:
            pdf.set_font("Helvetica", "", 10)
            for linha in conteudo:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(0, 6, _lat1(" | ".join(linha)))
            pdf.ln(1)

    return bytes(pdf.output())


# ── Excel (.xlsx) ───────────────────────────────────────────────


def gerar_xlsx(titulo: str, markdown: str) -> bytes:
    """Gera uma planilha: usa a primeira tabela do conteúdo, ou uma linha por item."""
    wb = Workbook()
    ws = wb.active
    ws.title = (titulo or "Planilha")[:31]
    blocos = parse_markdown(markdown)

    tabela = next((c for t, c in blocos if t == "tabela" and c), None)
    if tabela:
        for r_idx, linha in enumerate(tabela):
            ws.append(linha)
            if r_idx == 0:
                for cell in ws[ws.max_row]:
                    cell.font = XLFont(bold=True)
    else:
        for tipo, conteudo in blocos:
            if tipo in ("h1", "h2", "h3", "paragrafo"):
                ws.append([_sem_negrito(conteudo)])
            elif tipo == "lista":
                for item in conteudo:
                    ws.append([_sem_negrito(item)])

    for coluna in ws.columns:
        largura = max((len(str(c.value)) for c in coluna if c.value), default=10)
        ws.column_dimensions[coluna[0].column_letter].width = min(largura + 4, 80)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
